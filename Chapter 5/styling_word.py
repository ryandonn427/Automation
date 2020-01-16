import docx
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

document = docx.Document()

p = document.add_paragraph('This shows different kinds of emphasis')
p.add_run('bold').bold = True
p.add_run(', ')
p.add_run('italics').italic=True
p.add_run(' and ')
p.add_run('underline').underline=True
p.add_run('.')

document.add_paragraph('a few',style='List Bullet')
document.add_paragraph('bullet',style = 'List Bullet')
document.add_paragraph('points',style = 'List Bullet')
document.add_paragraph('Or numbered', style = 'List Number')
document.add_paragraph('that will', style = 'List Number')
document.add_paragraph('that keep', style = 'List Number')
document.add_paragraph('count',style = 'List Number')
document.add_paragraph('And finished with a quote', style = 'Quote')

p = document.add_paragraph('This paragraph will have manual styling and right alignment')
p.runs[0].font.name = 'Arial'
p.runs[0].font.size=Pt(25)
p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
document.save('word-report-style.docx')