import docx 
from datetime import datetime

context = {
    'date':datetime.now(),
    'movies':['Casablanca','The Sound of Music','Vertigo'],
    'total_minutes':404,
}

document = docx.Document()
document.add_heading('Movies Report',0)

paragraph = document.add_paragraph('Date: ')
paragraph.add_run(str(context['date'])).italic = True

paragraph = document.add_paragraph('Movies seen in the last 30 days:')
paragraph.add_run(str(len(context['movies']))).italic = True

for movie in context['movies']:
    document.add_paragraph(movie,style = 'List Bullet')

paragraph = document.add_paragraph('Total minutes: ')
paragraph.add_run(str(context['total_minutes'])).italic = True
document.save('word-report.docx')
