import docx
from docx.shared import Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
document = docx.Document()

document.add_paragraph('This is a document that includes a picture taken in Dublin')
image = document.add_picture('photo-dublin-a1.jpg')

image.width = Cm(14)
image.height = Cm(10)

paragraph = document.paragraphs[-1]
paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
paragraph.add_run().add_break()
paragraph.add_run('A picture of Dublin')

document.add_paragraph('Keep adding text after the image')
document.save('report.docx')